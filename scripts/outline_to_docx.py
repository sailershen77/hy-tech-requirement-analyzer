#!/usr/bin/env python3
"""把"技术部分目录结构"（Markdown 层级列表）转换为 docx 章节骨架。

层级与 Word 内置样式映射（中文 Word 中显示为"标题 1/2/3"）：
- 一级目录 "1. 实施方案〔20 分〕"  → Heading 1（标题 1）
- 二级目录 "1.1 项目概述与理解"     → Heading 2（标题 2）
- 三级要点 "- 建设背景与评估目标"   → Heading 3（标题 3）

用法:
    python outline_to_docx.py <input.md> <output.docx> [--title "文档标题"]

输入 markdown 逐行匹配，其余行（标题/表格/引用/正文）自动跳过：
- 无缩进 "1. xxx"（编号后须有空格）   → 标题 1
- 缩进 "1.1 xxx"                     → 标题 2
- 缩进 "1.1.1 xxx"                   → 标题 3
- 缩进 "- xxx" / "* xxx"             → 标题 3（去掉列表符号）

docx 输出时**剥离标题行开头的数字序列号**（如 "1."、"1.1"、"1.1.1"），
只保留标题文字、分值标注〔N 分〕与 Word 标题样式；编号应由 Word
多级列表/自动编号功能按标题样式自动生成。

排版约定：
- 每个标题行（Title / 标题 1/2/3）下面**追加一个空行**（Normal 样式的空段落），
  便于在骨架中直接填写章节内容
- 生成后**删除文档中未使用的多余样式**，只保留 Normal、Title、
  Heading 1/2/3（即"标题 1/2/3"），样式面板干净不干扰后续排版

依赖：python-docx
"""
import argparse
import re
import sys

from docx import Document

# 文档实际使用的样式（Normal 为默认正文样式，空行也用它，必须保留）
KEEP_STYLES = {"Normal", "Title", "Heading 1", "Heading 2", "Heading 3"}

# 无缩进一级有序列表（编号后必须有空格，避免误匹配 "1.1"）
H1_RE = re.compile(r"^(\d+)[.、]\s+(.+)$")
# 缩进的三级编号（先于二级检查）
H3_NUM_RE = re.compile(r"^\s+(\d+\.\d+\.\d+)\s+(.+)$")
# 缩进的二级编号
H2_RE = re.compile(r"^\s+(\d+\.\d+)\s+(.+)$")
# 缩进的列表要点
H3_DASH_RE = re.compile(r"^\s+[-*•]\s+(.+)$")

# 代码块围栏（跳过其中的内容，避免误提取示例代码）
FENCE_RE = re.compile(r"^```")


def clean_text(text):
    """去掉 markdown 内联标记（加粗等），保留分值标注。"""
    return text.replace("**", "").strip()


def strip_numbering(text):
    """剥离标题行开头的数字序列号（如 "1. "、"1.1 "、"1.1.1 "），
    保留标题文字与分值标注〔N 分〕。"""
    m = re.match(r"^\d+(?:\.\d+)*[.、]?\s*(.+)$", text)
    return m.group(1).strip() if m else text


def parse_markdown(lines):
    """返回 [(level, text), ...]，level ∈ {1, 2, 3}。"""
    items = []
    in_fence = False
    for raw in lines:
        line = raw.rstrip("\n")
        if FENCE_RE.match(line.strip()):
            in_fence = not in_fence
            continue
        if in_fence or not line.strip():
            continue

        m = H1_RE.match(line)
        if m and not line[0].isspace():
            items.append((1, strip_numbering(clean_text(f"{m.group(1)}. {m.group(2)}"))))
            continue
        m = H3_NUM_RE.match(line)
        if m:
            items.append((3, strip_numbering(clean_text(f"{m.group(1)} {m.group(2)}"))))
            continue
        m = H2_RE.match(line)
        if m:
            items.append((2, strip_numbering(clean_text(f"{m.group(1)} {m.group(2)}"))))
            continue
        m = H3_DASH_RE.match(line)
        if m:
            items.append((3, clean_text(m.group(1))))
    return items


def prune_unused_styles(doc):
    """删除文档中未使用的多余样式定义，只保留 KEEP_STYLES 中的样式。"""
    removed = 0
    for style in list(doc.styles):
        if style.name in KEEP_STYLES:
            continue
        try:
            style.delete()
            removed += 1
        except Exception:
            pass  # 个别内置样式不可删除时跳过
    return removed


def main():
    ap = argparse.ArgumentParser(
        description="技术部分目录结构 Markdown → docx（标题 1/2/3 样式）")
    ap.add_argument("input_md", help="输入 markdown 文件（含目录结构层级列表）")
    ap.add_argument("output_docx", help="输出 docx 文件路径")
    ap.add_argument("--title", default="技术部分目录结构",
                    help="docx 文档标题（Title 样式），默认“技术部分目录结构”")
    args = ap.parse_args()

    with open(args.input_md, encoding="utf-8") as f:
        items = parse_markdown(f.readlines())

    if not items:
        sys.exit("ERROR: 未在输入 markdown 中识别到目录层级行（一级 '1. '/二级 '1.1 '/三级 '- '）")

    doc = Document()
    doc.add_paragraph(clean_text(args.title), style="Title")
    doc.add_paragraph()  # 标题下空行
    for level, text in items:
        doc.add_paragraph(text, style=f"Heading {level}")
        doc.add_paragraph()  # 每个标题行下空行
    removed = prune_unused_styles(doc)
    doc.save(args.output_docx)

    counts = {1: 0, 2: 0, 3: 0}
    for lv, _ in items:
        counts[lv] += 1
    print(f"OK: {args.output_docx}")
    print(f"  标题 1: {counts[1]} 项 | 标题 2: {counts[2]} 项 | 标题 3: {counts[3]} 项")
    print(f"  每个标题下已加空行 | 已删除未使用样式 {removed} 个")


if __name__ == "__main__":
    main()
